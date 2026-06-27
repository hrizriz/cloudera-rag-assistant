from __future__ import annotations

import hashlib
import json
import zipfile
from io import BytesIO
from pathlib import Path

from cloudera_llm.config import ROOT_DIR, get_config
from cloudera_llm.ingestion.metadata import parse_local_metadata
from cloudera_llm.ingestion.scraper import ScrapedPage


SUPPORTED_EXTENSIONS = {".docx", ".doc", ".pdf", ".xlsx", ".xlsm"}


def load_local_documents(data_dir: Path | None = None) -> list[ScrapedPage]:
    config = get_config()
    root = data_dir or (ROOT_DIR / config.ingestion.local_data_dir)
    extensions = {ext.lower() for ext in config.ingestion.local_extensions}

    pages: list[ScrapedPage] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name == ".gitkeep":
            continue
        if "raw" in path.parts or "chroma" in path.parts:
            continue

        suffix = path.suffix.lower()
        if suffix == ".zip":
            pages.extend(_load_zip(path, extensions))
            continue
        if suffix not in extensions:
            continue
        page = _load_file(path)
        if page is not None:
            pages.append(page)

    return _dedupe_pages(pages)


def _dedupe_pages(pages: list[ScrapedPage]) -> list[ScrapedPage]:
    seen_hashes: set[str] = set()
    unique: list[ScrapedPage] = []
    skipped = 0

    for page in pages:
        content_hash = hashlib.sha1(page.text.encode("utf-8")).hexdigest()
        if content_hash in seen_hashes:
            skipped += 1
            continue
        seen_hashes.add(content_hash)
        unique.append(page)

    if skipped:
        print(f"[local] skipped {skipped} duplicate documents (same content)")
    return unique


def save_local_documents(pages: list[ScrapedPage]) -> int:
    from cloudera_llm.config import raw_data_path

    saved = 0
    for page in pages:
        slug = _safe_slug(page.title or Path(page.url).stem)
        digest = hashlib.sha1(page.url.encode("utf-8")).hexdigest()[:10]
        payload = {
            "url": page.url,
            "title": page.title,
            "text": page.text,
            "source_type": page.source_type,
            "product": page.product,
            "product_name": page.product_name,
            "version": page.version,
            "service": page.service,
            "doc_type": page.doc_type,
        }
        target = raw_data_path() / f"local_{slug}_{digest}.json"
        target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        saved += 1
    return saved


def _load_zip(path: Path, extensions: set[str]) -> list[ScrapedPage]:
    pages: list[ScrapedPage] = []
    try:
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                inner_suffix = Path(name).suffix.lower()
                if inner_suffix not in extensions:
                    continue
                with archive.open(name) as handle:
                    data = handle.read()
                inner_path = Path(name)
                page = _load_bytes(data, inner_path, source_root=path)
                if page is not None:
                    pages.append(page)
    except (zipfile.BadZipFile, OSError) as exc:
        print(f"[skip] {path}: {exc}")
    return pages


def _load_file(path: Path) -> ScrapedPage | None:
    try:
        data = path.read_bytes()
    except OSError as exc:
        print(f"[skip] {path}: {exc}")
        return None
    return _load_bytes(data, path, source_root=path.parent)


def _load_bytes(data: bytes, path: Path, *, source_root: Path) -> ScrapedPage | None:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            text = _read_docx(data)
        elif suffix == ".pdf":
            text = _read_pdf(data)
        elif suffix in {".xlsx", ".xlsm"}:
            text = _read_xlsx(data)
        elif suffix == ".doc":
            text = _read_doc_fallback(path.name, data)
        else:
            return None
    except Exception as exc:
        print(f"[skip] {path}: {exc}")
        return None

    text = _normalize_text(text)
    if len(text) < 80:
        print(f"[skip] {path}: extracted text too short")
        return None

    title = path.stem.replace("_", " ").strip() or path.name
    if source_root.is_file():
        source_url = f"local://{source_root.resolve().as_posix()}#{path.as_posix()}"
    else:
        source_url = f"local://{path.resolve().as_posix()}"

    meta = parse_local_metadata(title, source_url)
    return ScrapedPage(
        url=source_url,
        title=title,
        text=text,
        source_type="local",
        product=meta.product,
        product_name=meta.product_name,
        version=meta.version,
        service=meta.service,
        doc_type=meta.doc_type,
        is_support_matrix=False,
    )


def _read_docx(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line:
            parts.append(line)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _read_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() for value in row if value not in (None, "")]
            if cells:
                parts.append(" | ".join(cells))
    workbook.close()
    return "\n".join(parts)


def _read_doc_fallback(filename: str, data: bytes) -> str:
    # Legacy .doc is not reliably supported without extra native tools.
    del data
    raise ValueError(f"legacy .doc not supported ({filename}); convert to .docx")


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned = [line for line in lines if line]
    return "\n".join(cleaned)


def _safe_slug(value: str) -> str:
    slug = "".join(ch if ch.isalnum() else "_" for ch in value.lower()).strip("_")
    return slug[:80] or "document"
